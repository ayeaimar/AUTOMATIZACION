from docx import Document
import pandas as pd

def parse_docx_tasks(file_path: str):
    """
    Extrae la estructura del documento.
    Si el documento posee tablas (Proceso | Tarea), las lee.
    De lo contrario, lee párrafos estructurados.
    """
    doc = Document(file_path)
    tasks = []
    
    # Intenta extraer de tablas
    for table in doc.tables:
        for row in table.rows[1:]: # Ignorar encabezado
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                proceso = cells[0]
                tarea = cells[1]
                tasks.append({"Proceso": proceso, "Tarea": tarea})
                
    # Fallback a párrafos si no hay tablas
    if not tasks:
        current_proceso = "General"
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if p.style.name.startswith('Heading'):
                current_proceso = text
            else:
                tasks.append({"Proceso": current_proceso, "Tarea": text})
                
    return tasks

def compare_documents(old_file: str, new_file: str):
    old_tasks = parse_docx_tasks(old_file)
    new_tasks = parse_docx_tasks(new_file)
    
    old_map = {t["Tarea"]: t["Proceso"] for t in old_tasks}
    new_map = {t["Tarea"]: t["Proceso"] for t in new_tasks}
    
    results = []
    
    # Tareas en el nuevo documento
    for tarea, proceso in new_map.items():
        if tarea in old_map:
            # Si existía pero tuvo cambios internos
            results.append({
                "Proceso": proceso,
                "Tarea": tarea,
                "Estado": "MODIFICADA",
                "Detalle": "Contenido actualizado"
            })
        else:
            results.append({
                "Proceso": proceso,
                "Tarea": tarea,
                "Estado": "AGREGADA",
                "Detalle": "Nueva tarea"
            })
            
    # Tareas eliminadas
    for tarea, proceso in old_map.items():
        if tarea not in new_map:
            results.append({
                "Proceso": proceso,
                "Tarea": tarea,
                "Estado": "ELIMINADA",
                "Detalle": "Tarea removida"
            })
            
    df = pd.DataFrame(results)
    
    counts = {
        "agregadas": int((df["Estado"] == "AGREGADA").sum()) if not df.empty else 0,
        "modificadas": int((df["Estado"] == "MODIFICADA").sum()) if not df.empty else 0,
        "eliminadas": int((df["Estado"] == "ELIMINADA").sum()) if not df.empty else 0
    }
    
    return df, counts
