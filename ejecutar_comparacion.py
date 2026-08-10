from datetime import datetime
import os
import sys
from comparador import comparar
from docx import Document
from extractor import extraer_tabla
import pandas as pd

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
carpeta_anterior = os.path.join(BASE_DIR, "Historial_Viejo")
carpeta_nuevo = os.path.join(BASE_DIR, "Historial_Nuevo")
carpeta_resultados = os.path.join(BASE_DIR, "Resultados")

os.makedirs(carpeta_resultados, exist_ok=True)

if len(sys.argv) < 3:
    raise Exception(
        "Debe recibir los nombres de los archivos: python ejecutar_comparacion.py <archivo_viejo.docx> <archivo_nuevo.docx>"
    )

nombre_archivo_viejo = sys.argv[1]
nombre_archivo_nuevo = sys.argv[2]

archivo_anterior = os.path.join(carpeta_anterior, nombre_archivo_viejo)
archivo_nuevo = os.path.join(carpeta_nuevo, nombre_archivo_nuevo)

if not os.path.exists(archivo_anterior):
    raise Exception(f"No existe el archivo viejo: {archivo_anterior}")

if not os.path.exists(archivo_nuevo):
    raise Exception(f"No existe el archivo nuevo: {archivo_nuevo}")

# Extracción de tablas a DataFrames
df_anterior = extraer_tabla(archivo_anterior)
df_nuevo = extraer_tabla(archivo_nuevo)

# Ejecución de la comparación
resultado = comparar(df_anterior, df_nuevo)

fecha = datetime.now().strftime("%Y%m%d_%H%M%S")

# 1. Guardar resultados en Excel local (si es DataFrame)
excel_path = os.path.join(carpeta_resultados, f"Comparacion_{fecha}.xlsx")
if isinstance(resultado, pd.DataFrame):
    df_resultado = resultado
else:
    df_resultado = pd.DataFrame(resultado)

df_resultado.to_excel(excel_path, index=False)

# 2. Conteo de métricas
agregadas = (
    len(df_resultado[df_resultado["Estado"] == "AGREGADA"])
    if not df_resultado.empty
    else 0
)
modificadas = (
    len(df_resultado[df_resultado["Estado"] == "MODIFICADA"])
    if not df_resultado.empty
    else 0
)
eliminadas = (
    len(df_resultado[df_resultado["Estado"] == "ELIMINADA"])
    if not df_resultado.empty
    else 0
)

# 3. Construcción del texto descriptivo
detalle_completo = ""
if not df_resultado.empty:
    for _, fila in df_resultado.iterrows():
        detalle_completo += (
            f"[{fila['Estado']}] Proceso: {fila['Proceso']} | "
            f"Tarea: {fila['Tarea']}\n"
            f"Detalle: {fila['Detalle']}\n\n"
        )

# 4. Generación de informe en Word local
doc = Document()
doc.add_heading("Informe de Comparación de Procedimientos", level=1)
doc.add_paragraph(f"Documento anterior: {os.path.basename(archivo_anterior)}")
doc.add_paragraph(f"Documento nuevo: {os.path.basename(archivo_nuevo)}")
doc.add_paragraph(f"Agregadas: {agregadas}")
doc.add_paragraph(f"Modificadas: {modificadas}")
doc.add_paragraph(f"Eliminadas: {eliminadas}")

doc.add_heading("Cambios detectados", level=2)
if detalle_completo.strip():
    doc.add_paragraph(detalle_completo)
else:
    doc.add_paragraph("No se detectaron cambios entre ambas versiones.")

docx_path = os.path.join(carpeta_resultados, f"Informe_Comparacion_{fecha}.docx")
doc.save(docx_path)

# Salida limpia por consola
print(
    f"VIEJO: {os.path.basename(archivo_anterior)}\n"
    f"NUEVO: {os.path.basename(archivo_nuevo)}\n"
    f"Métricas -> Agregadas: {agregadas} | Modificadas: {modificadas} | Eliminadas: {eliminadas}\n"
)
